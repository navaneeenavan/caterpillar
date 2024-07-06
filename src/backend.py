from flask import Flask, jsonify, request,send_file
import google.generativeai as genai
import re
import json
from pymongo import MongoClient
from flask_cors import CORS
from pymongo import MongoClient
from io import BytesIO
from docx.shared import Inches
import base64
import requests

app = Flask(__name__)
CORS(app)
cors = CORS(app, resources={r"/api/*": {"origins": "http://localhost:5000"}})
# Configure Gemini API
genai.configure(api_key="AIzaSyDXvkJJYAbSDag8EnppGOpzpWmQ8aPUYnY")

# MongoDB connection URI
mongo_uri = "mongodb+srv://admin:admin@cluster0.9fknthw.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"

# Connect to MongoDB server
client = MongoClient(mongo_uri)

# Select the database
db = client['inspectionsDB']  # Replace with your database name

# Select the collection
collection = db['inspections']  # Replace with your collection name


@app.route('/updateform/<truck_id>', methods=['POST'])
def update_form(truck_id):
    data = request.json
    try:
        new_record = {
            "battery": {
                "condition": {
                    "any_damage": "",
                    "damage_image": None
                },
                "leak_in_battery": False,
                "make": "",
                "replacement_date": "",
                "rust_in_battery": False,
                "voltage": "",
                "water_level": ""
            },
            "brakes_and_exterior": {
                "attached_images": [],
                "brake_fluid_level": 0,
                "brake_overall_summary": "",
                "dent": False,
                "emergency_brake": False,
                "front_brake_condition": "",
                "oil_leak_in_suspension": False,
                "rear_brake_condition": "",
                "rust": False
            },
            "cat_customer_id": "",
            "customer_name": "",
            "engine": {
                "attached_images": [],
                "brake_fluid_color": "",
                "brake_fluid_condition": "",
                "dent_damage": False,
                "engine_oil_color": "",
                "engine_oil_condition": "",
                "oil_leak_in_engine": "",
                "overall_summary": "",
                "rust_damage": False
            },
            "geo_coordinates": "",
            "inspection_datetime": "",
            "inspection_employee_id": "",
            "inspection_id": "",
            "inspector_name": "",
            "inspector_signature": "",
            "location": "",
            "service_meter_hours": 0,
            "tires": {
                "left_front_condition": "",
                "left_front_pressure": 0,
                "left_rear_condition": "",
                "left_rear_pressure": 0,
                "right_front_condition": "",
                "right_front_pressure": 0,
                "right_rear_condition": "",
                "right_rear_pressure": 0
            },
            "truck_model": "",
            "truck_serial_number": ""
        }

        # Update the new record with the values from the form data
        for key in data:
            if key in new_record and isinstance(new_record[key], dict):
                new_record[key].update(data[key])
            else:
                new_record[key] = data[key]

        # Insert the new record into the collection
        result = collection.insert_one(new_record)

        return jsonify({"message": "Form submitted successfully!", "id": str(result.inserted_id)}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/inspections', methods=['GET'])

def get_inspections():
    """
    Fetch all inspection records from the database.
    
    :return: JSON list of all inspection records.
    """
    inspections = collection.find()
    results = []
    for inspection in inspections:
        # Convert MongoDB ObjectId to string for JSON serialization
        inspection['_id'] = str(inspection['_id'])
        results.append(inspection)
    return jsonify(results)

# Helper function to update a specific section
def update_section(truck_serial_number, section_name, section_data):
    """
    Update a specific section for a truck based on its serial number.
    
    :param truck_serial_number: The serial number of the truck.
    :param section_name: The name of the section to update.
    :param section_data: The data to update in the section.
    :return: JSON response with the status of the update.
    """
    # Find the truck inspection record by serial number
    inspection = collection.find_one({"truck_serial_number": truck_serial_number})
    
    if not inspection:
        return jsonify({"error": f"Truck with serial number {truck_serial_number} not found"}), 404
    
    # Prepare the update query
    update_query = {}
    for key, value in section_data.items():
        if key in inspection[section_name]:
            update_query[f'{section_name}.{key}'] = value
    
    # Perform the update
    if update_query:
        result = collection.update_one(
            {"truck_serial_number": truck_serial_number},
            {"$set": update_query}
        )
        
        if result.modified_count > 0:
            return jsonify({"message": f"{section_name.capitalize()} updated successfully"})
        else:
            return jsonify({"message": f"No changes made to the {section_name} data"})
    
    return jsonify({"error": f"No valid {section_name} data provided"}), 400

# Route to update tires
@app.route('/api/truck-inspections/update-tires/<truck_serial_number>', methods=['POST'])
def update_tires(truck_serial_number):
    """
    Update the tires information for a specific truck.
    
    :param truck_serial_number: The serial number of the truck.
    :return: JSON response with the status of the update.
    """
    data = request.json
    print(data)
    tires_data = data.get('tire_info', {})
    return update_section(truck_serial_number, 'tires', tires_data)

# Route to update battery
@app.route('/api/truck-inspections/update-battery/<truck_serial_number>', methods=['POST'])
def update_battery(truck_serial_number):
    """
    Update the battery information for a specific truck.
    
    :param truck_serial_number: The serial number of the truck.
    :return: JSON response with the status of the update.
    """
    data = request.json
    battery_data = data.get('battery', {})
    return update_section(truck_serial_number, 'battery', battery_data)

# Route to update engine
@app.route('/api/truck-inspections/update-engine/<truck_serial_number>', methods=['POST'])
def update_engine(truck_serial_number):
    """
    Update the engine information for a specific truck.
    
    :param truck_serial_number: The serial number of the truck.
    :return: JSON response with the status of the update.
    """
    
    data = request.json
    
    engine_data = data.get('engine', {})
    return update_section(truck_serial_number, 'engine', engine_data)

# Route to update brakes and exterior
@app.route('/api/truck-inspections/update-brakes-and-exterior/<truck_serial_number>', methods=['POST'])
def update_brakes_and_exterior(truck_serial_number):
    """
    Update the brakes and exterior information for a specific truck.
    
    :param truck_serial_number: The serial number of the truck.
    :return: JSON response with the status of the update.
    """
    data = request.json
    print(data)
    brakes_and_exterior_data = data.get('brakes_and_exterior', {})
    
    return update_section(truck_serial_number, 'brakes_and_exterior', brakes_and_exterior_data)

def generate_battery_info(text):
    """
    Generate battery information from text using Gemini API.
    
    :param text: The input text describing the battery details.
    :return: Parsed JSON object with battery details.
    """
    try:
        prompt = f"""
        You are an expert in interpreting technical details from textual descriptions.
        Given the following text, extract the information related to battery details and fill out the corresponding fields.
        If any information is not available in the text, leave it empty. Here are the fields you need to fill:
        - Battery make
        - Battery replacement date
        - Battery voltage
        - Battery water level (Good, Ok, Low)
        - Condition of battery (Yes/No)
        - Any Leak (Yes/No)
        - Any Rust (Yes/No)

        Text: {text}
        for date add in this format yyyy-MM-dd
        Provide the output as a JSON object with the following format:
        {{
            "truck_serial_number" : " ",
            "battery_make": "",
            "battery_replacement_date": "",
            "battery_voltage": "",
            "battery_water_level": "",
            "condition_of_battery": "",
            "any_leak": "",
            "any_rust": "",
            "overall_summary": "",
        }}
        Only provide the JSON object for easier parsing and formatting, without any additional formatting or markdown.
        """

        model = genai.GenerativeModel(model_name="models/gemini-1.5-flash")
        response = model.generate_content([prompt])

        if response.candidates:
            response_text = response.candidates[0].content.parts[0].text.strip()
            print(f"Response from Gemini API: {response_text}")
            
            if not response_text:
                return {'error': 'Empty response from API'}
            
            data = split_and_load_ejson(response_text)
            
            if data is None:
                print("Attempting to parse response as plain JSON.")
                data = clean_response(response_text)
            
            if data is None:
                return {'error': 'Failed to extract JSON from API response'}
            
            return data
        else:
            return {'error': 'No candidates in API response'}
    
    except Exception as e:
        print(f"Error during API call or processing: {str(e)}")
        return {'error': str(e)}

def split_and_load_ejson(text):
    """
    Extract JSON content from a text block containing JSON within triple backticks.
    
    :param text: The input text containing JSON content.
    :return: Parsed JSON object or None if extraction fails.
    """
    pattern = r"json(.*?)"
    match = re.search(pattern, text, re.DOTALL)
    
    if match:
        json_content = match.group(1).strip()
        try:
            data = json.loads(json_content)
            return data
        except json.JSONDecodeError as e:
            print(f"Error: Invalid JSON format in extracted content: {str(e)}")
            return None
    else:
        print("Error: No valid JSON block found using triple backticks")
        return None

def clean_response(response_text):
    """
    Extract a valid JSON object from the response text.
    
    :param response_text: The raw text response from the API.
    :return: Parsed JSON object or None if extraction fails.
    """
    try:
        json_data = json.loads(response_text)
        return json_data
    except json.JSONDecodeError:
        pass

    json_pattern = re.compile(r'({.*})', re.DOTALL)
    match = json_pattern.search(response_text)
    if match:
        json_content = match.group(1)
        try:
            json_data = json.loads(json_content)
            return json_data
        except json.JSONDecodeError as e:
            print(f"Error: Invalid JSON format in cleaned content: {str(e)}")
            return None
    
    print("Error: No valid JSON object found in response")
    return None

# Route to generate battery information from text
@app.route('/generate_battery_info', methods=['POST'])
def generate_battery_info_route():
    """
    Extract battery information from a given text using the Gemini API.
    
    :return: JSON object containing extracted battery information.
    """
    data = request.get_json()
    text = data.get('text')

    if not text:
        return jsonify({'error': 'Text is required'}), 400

    battery_info = generate_battery_info(text)
    
    if battery_info is None or 'error' in battery_info:
        return jsonify(battery_info), 500

    return jsonify({'battery_info': battery_info}), 200



def generate_engine_info(text):
    """
    Generate engine information from text using Gemini API.
    
    :param text: The input text describing the engine details.
    :return: Parsed JSON object with engine details.
    """
    try:
        prompt = f"""
        You are an expert in interpreting technical details from textual descriptions.
        Given the following text, extract the information related to engine details and fill out the corresponding fields.
        If any information is not available in the text, leave it empty. Here are the fields you need to fill:
        - Truck serial number
        - Attached images (URLs)
        - Brake fluid color (clean, brown , black)
        - Brake fluid condition (Good , Bad)
        - Dent damage (true/false)
        - Engine oil color (clean, brown , black)
        - Engine oil condition
        - Oil leak in engine (Yes/No)
        - Rust damage (true/false)
        - Overall summary

        Text: {text}
        Provide the output as a JSON object with the following format:
        {{
            "truck_serial_number": "",
            "brake_fluid_color": "",
            "brake_fluid_condition": "",
            "dent_damage": "",
            "engine_oil_color": "",
            "engine_oil_condition": "",
            "oil_leak_in_engine": "",
            "rust_damage": "",
            "overall_summary": ""
        }}
        Only provide the JSON object for easier parsing and formatting, without any additional formatting or markdown.
        """

        model = genai.GenerativeModel(model_name="models/gemini-1.5-flash")
        response = model.generate_content([prompt])

        if response.candidates:
            response_text = response.candidates[0].content.parts[0].text.strip()
            print(f"Response from Gemini API: {response_text}")
            
            if not response_text:
                return {'error': 'Empty response from API'}
            
            data = split_and_load_ejson(response_text)
            
            if data is None:
                print("Attempting to parse response as plain JSON.")
                data = clean_response(response_text)
            
            if data is None:
                return {'error': 'Failed to extract JSON from API response'}
            
            return data
        else:
            return {'error': 'No candidates in API response'}
    
    except Exception as e:
        print(f"Error during API call or processing: {str(e)}")
        return {'error': str(e)}
    
@app.route('/generate_engine_info', methods=['POST'])
def generate_engine_info_route():
    """
    Extract engine information from a given text using the Gemini API.
    
    :return: JSON object containing extracted engine information.
    """
    data = request.get_json()
    text = data.get('text')

    if not text:
        return jsonify({'error': 'Text is required'}), 400

    engine_info = generate_engine_info(text)
    
    if engine_info is None or 'error' in engine_info:
        return jsonify(engine_info), 500

    return jsonify({'engine_info': engine_info}), 200

def generate_brakes_and_exterior_info(text):
    """
    Generate brakes and exterior information from text using Gemini API.
    
    :param text: The input text describing the brakes and exterior details.
    :return: Parsed JSON object with brakes and exterior details.
    """
    try:
        prompt = f"""
        You are an expert in interpreting technical details from textual descriptions.
        Given the following text, extract the information related to brakes and exterior details and fill out the corresponding fields.
        If any information is not available in the text, leave it empty. Here are the fields you need to fill:
        - Truck serial number
        - Attached images (URLs) (attach random urls)
        - Brake fluid level (not a string)
        - Brake overall summary
        - Dent (true/false)
        - Emergency brake (true/false)
        - Front brake condition (Good, Fair, Needs Replacement)
        - Oil leak in suspension (Yes/No)
        - Rear brake condition (Good, Fair, Needs Replacement)
        - Rust (true/false)

        Text: {text}
        Provide the output as a JSON object with the following format:
        {{
            "truck_serial_number": "",
            "attached_images": [],
            "brake_fluid_level": 10,
            "brake_overall_summary": "",
            "dent": false,
            "emergency_brake": "",
            "front_brake_condition": "",
            "oil_leak_in_suspension": false,
            "rear_brake_condition": "",
            "rust": false
        }}
        Even though you can't find it in the text input just make sure that you are filing ever feilds on the basis of your assumptions even with the summary and fluid level too
        Only provide the JSON object for easier parsing and formatting, without any additional formatting or markdown even though somethough someof the values are not given just give a assumed value based on teh other values.
        """

        model = genai.GenerativeModel(model_name="models/gemini-1.5-flash")
        response = model.generate_content([prompt])

        if response.candidates:
            response_text = response.candidates[0].content.parts[0].text.strip()
            print(f"Response from Gemini API: {response_text}")
            
            if not response_text:
                return {'error': 'Empty response from API'}
            
            data = split_and_load_ejson(response_text)
            
            if data is None:
                print("Attempting to parse response as plain JSON.")
                data = clean_response(response_text)
            
            if data is None:
                return {'error': 'Failed to extract JSON from API response'}
            
            return data
        else:
            return {'error': 'No candidates in API response'}
    
    except Exception as e:
        print(f"Error during API call or processing: {str(e)}")
        return {'error': str(e)}

@app.route('/generate_brakes_and_exterior_info', methods=['POST'])
def generate_brakes_and_exterior_info_route():
    """
    Extract brakes and exterior information from a given text using the Gemini API.
    
    :return: JSON object containing extracted brakes and exterior information.
    """
    data = request.get_json()
    text = data.get('text')

    if not text:
        return jsonify({'error': 'Text is required'}), 400

    brakes_and_exterior_info = generate_brakes_and_exterior_info(text)
    
    if brakes_and_exterior_info is None or 'error' in brakes_and_exterior_info:
        return jsonify(brakes_and_exterior_info), 500

    return jsonify({'brakes_and_exterior_info': brakes_and_exterior_info}), 200






def generate_tire_info(text):
    """
    Generate tire information from text using Gemini API.
    
    :param text: The input text describing the tire details.
    :return: Parsed JSON object with tire details.
    """
    try:
        prompt = f"""
        You are an expert in interpreting technical details from textual descriptions.
        Given the following text, extract the information related to tire details and fill out the corresponding fields.
        If any information is not available in the text, leave it empty. Here are the fields you need to fill:
        - Truck serial number
        - Tire pressure (Left Front, Right Front, Left Rear, Right Rear) (number)
        - Tire condition (Left Front, Right Front, Left Rear, Right Rear) (number)
        - Tire images (URLs)

        Text: {text}
        Provide the output as a JSON object with the following format:
        {{
            "truck_serial_number": "",
            "tires": {{
                "left_front_pressure": "",
                "left_front_condition": "",
                "right_front_pressure": "",
                "right_front_condition": "",
                "left_rear_pressure": "",
                "left_rear_condition": "",
                "right_rear_pressure": "",
                "right_rear_condition": ""
            }},
            "tire_images": [],
            "overall_summary": ""
        }}
        if the value was not specified in the propmt use the sentiment and avaiblable information of the promt and fill the values automatically as your assuumptions even with the tire pressure too
        Only provide the JSON object for easier parsing and formatting, without any additional formatting or markdown.
        """

        model = genai.GenerativeModel(model_name="models/gemini-1.5-flash")
        response = model.generate_content([prompt])

        if response.candidates:
            response_text = response.candidates[0].content.parts[0].text.strip()
            print(f"Response from Gemini API: {response_text}")
            
            if not response_text:
                return {'error': 'Empty response from API'}
            
            data = split_and_load_ejson(response_text)
            
            if data is None:
                print("Attempting to parse response as plain JSON.")
                data = clean_response(response_text)
            
            if data is None:
                return {'error': 'Failed to extract JSON from API response'}
            
            return data
        else:
            return {'error': 'No candidates in API response'}
    
    except Exception as e:
        print(f"Error during API call or processing: {str(e)}")
        return {'error': str(e)}

@app.route('/api/generate-tire-info', methods=['POST'])
def generate_tire_info_route():
    """
    Extract tire information from a given text using the Gemini API.
    
    :return: JSON object containing extracted tire information.
    """
    data = request.get_json()
    text = data.get('text')

    if not text:
        return jsonify({'error': 'Text is required'}), 400

    tire_info = generate_tire_info(text)
    
    if tire_info is None or 'error' in tire_info:
        return jsonify(tire_info), 500

    return jsonify({'tire_info': tire_info}), 200


def add_image_from_url(doc, url, width=Inches(1.25)):
    response = requests.get(url)
    if response.status_code == 200:
        img_path = "temp_image.jpg"
        with open(img_path, "wb") as f:
            f.write(response.content)
        doc.add_picture(img_path, width=width)

def generate_sample_report(data):
    doc = Document()
    doc.add_heading('Sample MongoDB Data Report', level=1).alignment = 1

    for record in data:
        doc.add_heading('Inspection Report', level=2)

        doc.add_heading('Truck Details', level=3)
        doc.add_paragraph(f"Serial Number: {record.get('truck_serial_number', '')}")
        doc.add_paragraph(f"Model: {record.get('truck_model', '')}")

        doc.add_heading('Inspection Details', level=3)
        doc.add_paragraph(f"Inspection ID: {record.get('inspection_id', '')}")
        doc.add_paragraph(f"Inspector Name: {record.get('inspector_name', '')}")
        doc.add_paragraph(f"Inspector Employee ID: {record.get('inspection_employee_id', '')}")
        doc.add_paragraph(f"Inspection Date & Time: {record.get('inspection_datetime', '')}")
        doc.add_paragraph(f"Location: {record.get('location', '')}")
        doc.add_paragraph(f"Geo Coordinates: {record.get('geo_coordinates', '')}")
        doc.add_paragraph(f"Service Meter Hours: {record.get('service_meter_hours', '')}")
        doc.add_paragraph(f"Customer Name: {record.get('customer_name', '')}")
        doc.add_paragraph(f"Customer ID: {record.get('cat_customer_id', '')}")

        doc.add_heading('Tires Details', level=3)
        tires = record.get('tires', {})
        doc.add_paragraph(f"Left Front Pressure: {tires.get('left_front_pressure', '')}")
        doc.add_paragraph(f"Right Front Pressure: {tires.get('right_front_pressure', '')}")
        doc.add_paragraph(f"Left Front Condition: {tires.get('left_front_condition', '')}")
        doc.add_paragraph(f"Right Front Condition: {tires.get('right_front_condition', '')}")
        doc.add_paragraph(f"Left Rear Pressure: {tires.get('left_rear_pressure', '')}")
        doc.add_paragraph(f"Right Rear Pressure: {tires.get('right_rear_pressure', '')}")
        doc.add_paragraph(f"Left Rear Condition: {tires.get('left_rear_condition', '')}")
        doc.add_paragraph(f"Right Rear Condition: {tires.get('right_rear_condition', '')}")

        doc.add_heading('Battery Details', level=3)
        battery = record.get('battery', {})
        doc.add_paragraph(f"Make: {battery.get('make', '')}")
        doc.add_paragraph(f"Replacement Date: {battery.get('replacement_date', '')}")
        doc.add_paragraph(f"Voltage: {battery.get('voltage', '')}")
        doc.add_paragraph(f"Water Level: {battery.get('water_level', '')}")
        condition = battery.get('condition', {})
        doc.add_paragraph(f"Any Damage: {condition.get('any_damage', '')}")
        doc.add_paragraph(f"Leak or Rust: {battery.get('leak_or_rust', '')}")

        if battery.get("image_url"):
            add_image_from_url(doc, battery["image_url"])
        if tires.get("image_url"):
            add_image_from_url(doc, tires["image_url"])

        doc.add_page_break()

    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    return byte_io

@app.route('/generate_report/<inspection_id>', methods=['GET'])
def generate_report(inspection_id):
    try:
        data = list(collection.find({"truck_serial_number": inspection_id}))
        print(data)
        if not data:
            return jsonify({"error": "No data found for truck serial number"}), 404
    except PyMongoError as e:
        return jsonify({"error": f"MongoDB error: {str(e)}"}), 500

    doc_io = generate_sample_report(data)
    return send_file(doc_io, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, attachment_filename=f"{inspection_id}_report.docx")


if __name__ == '__main__':
    app.run(debug=True)